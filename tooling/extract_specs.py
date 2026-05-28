"""Extract Korean spec docs (.docx, .xlsx) to plain-text files for analysis."""
import sys
import os
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from openpyxl import load_workbook

ROOT = Path("docs/명세서")
OUT = Path("docs/_extracted")
OUT.mkdir(parents=True, exist_ok=True)


def dump_docx(path: Path, out_path: Path):
    doc = Document(path)
    lines = []
    # Walk in document order: paragraphs and tables interleaved via doc.element.body
    body = doc.element.body
    from docx.oxml.ns import qn

    table_iter = iter(doc.tables)
    para_iter = iter(doc.paragraphs)
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)
    pi = 0
    ti = 0
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            if pi < len(paragraphs):
                p = paragraphs[pi]
                pi += 1
                txt = p.text.strip()
                if txt:
                    lines.append(txt)
        elif tag == "tbl":
            if ti < len(tables):
                t = tables[ti]
                ti += 1
                lines.append(f"\n--- TABLE ---")
                for row in t.rows:
                    cells = []
                    for c in row.cells:
                        cells.append(c.text.strip().replace("\n", " / "))
                    lines.append(" | ".join(cells))
                lines.append("--- /TABLE ---\n")
    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"WROTE {out_path} ({len(lines)} lines)")


def dump_xlsx(path: Path, out_path: Path):
    wb = load_workbook(path, data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"\n========== SHEET: {ws.title} ==========\n")
        for row in ws.iter_rows(values_only=True):
            cells = [("" if c is None else str(c)) for c in row]
            if any(c.strip() for c in cells):
                lines.append(" | ".join(cells))
    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"WROTE {out_path} ({len(lines)} lines)")


def main():
    for f in sorted(ROOT.iterdir()):
        if f.is_dir():
            continue
        stem = f.stem.replace(" ", "_").replace(".", "")
        if f.suffix.lower() == ".docx":
            dump_docx(f, OUT / f"{stem}.txt")
        elif f.suffix.lower() == ".xlsx":
            dump_xlsx(f, OUT / f"{stem}.txt")


if __name__ == "__main__":
    main()
